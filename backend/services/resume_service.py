"""
Resume Processing Service
Handles resume file upload, parsing, and text extraction
"""

import os
import json
import PyPDF2
import docx
from models import db, Resume, UserSkill
from services.skill_extractor import SkillExtractor

# Optional NLP imports
try:
    import spacy
    SPACY_AVAILABLE = True
except ImportError:
    SPACY_AVAILABLE = False

try:
    import nltk
    from nltk.tokenize import sent_tokenize, word_tokenize
    NLTK_AVAILABLE = True
    # Download required NLTK data
    try:
        nltk.data.find('tokenizers/punkt')
    except LookupError:
        try:
            nltk.download('punkt')
        except Exception as e:
            NLTK_AVAILABLE = False
except ImportError:
    NLTK_AVAILABLE = False

class ResumeService:
    UPLOAD_FOLDER = 'uploads/resumes'
    ALLOWED_EXTENSIONS = {'pdf', 'docx'}
    
    def __init__(self):
        if SPACY_AVAILABLE:
            try:
                self.nlp = spacy.load('en_core_web_sm')
            except Exception as e:
                self.nlp = None
        else:
            self.nlp = None
        self.skill_extractor = SkillExtractor()
        os.makedirs(self.UPLOAD_FOLDER, exist_ok=True)
    
    @staticmethod
    def extract_text_from_pdf(file_path):
        """Extract text from PDF file"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    text += page.extract_text()
        except Exception as e:
            return None, f"Error reading PDF: {str(e)}"
        
        return text, None
    
    @staticmethod
    def extract_text_from_docx(file_path):
        """Extract text from DOCX file"""
        text = ""
        try:
            doc = docx.Document(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
        except Exception as e:
            return None, f"Error reading DOCX: {str(e)}"
        
        return text, None
    
    def process_resume(self, user_id, file, filename):
        """
        Process uploaded resume file
        Returns: (resume, error)
        """
        # Determine file type
        file_ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
        
        if file_ext not in self.ALLOWED_EXTENSIONS:
            return None, f"Invalid file type. Allowed: {', '.join(self.ALLOWED_EXTENSIONS)}"
        
        # Save file
        file_path = os.path.join(self.UPLOAD_FOLDER, f"{user_id}_{filename}")
        file.save(file_path)
        
        # Extract text
        if file_ext == 'pdf':
            text, error = self.extract_text_from_pdf(file_path)
        else:  # docx
            text, error = self.extract_text_from_docx(file_path)
        
        if error:
            return None, error
        
        # Parse resume
        parsed_data = self._parse_resume_text(text)
        
        # Store in database
        resume = Resume.query.filter_by(user_id=user_id).first()
        if resume:
            os.remove(resume.file_path)  # Remove old file
            resume.filename = filename
            resume.file_path = file_path
            resume.file_type = file_ext
            resume.raw_text = text
            resume.parsed_data = json.dumps(parsed_data)
        else:
            resume = Resume(
                user_id=user_id,
                filename=filename,
                file_path=file_path,
                file_type=file_ext,
                raw_text=text,
                parsed_data=json.dumps(parsed_data)
            )
            db.session.add(resume)
        
        try:
            db.session.commit()
            
            # Auto-populate user skills from resume
            self._populate_skills_from_resume(user_id, parsed_data)
            
            return resume, None
        except Exception as e:
            db.session.rollback()
            return None, str(e)
    
    def _parse_resume_text(self, text):
        """Parse resume text and extract structured data"""
        # Use spaCy NER for entity extraction
        doc = self.nlp(text)
        
        skills = self.skill_extractor.extract_skills(text)
        
        # Extract education
        education = self._extract_education(text)
        
        # Extract experience years
        experience_years = self._extract_experience_years(text)
        
        return {
            'skills': skills,
            'education': education,
            'experience_years': experience_years,
            'raw_entities': [{'text': ent.text, 'label': ent.label_} for ent in doc.ents]
        }
    
    def _extract_education(self, text):
        """Extract education information"""
        degrees = ['Bachelor', 'Master', 'PhD', 'Diploma', 'B.Tech', 'B.E', 'M.Tech', 'MBA']
        education = []
        
        for degree in degrees:
            if degree.lower() in text.lower():
                education.append(degree)
        
        return education if education else ['Not specified']
    
    def _extract_experience_years(self, text):
        """Extract years of experience"""
        import re
        
        patterns = [
            r'(\d+)\s*(?:years?|yrs?)\s*(?:of\s*)?(?:experience|exp)',
            r'(?:experience|exp).*?(\d+)\s*(?:years?|yrs?)',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                try:
                    return float(match.group(1))
                except ValueError:
                    pass
        
        return 0
    
    def _populate_skills_from_resume(self, user_id, parsed_data):
        """Auto-populate user skills from parsed resume"""
        skills = parsed_data.get('skills', [])
        
        for skill in skills:
            # Add skill to user profile
            user_skill = UserSkill.query.filter_by(user_id=user_id, skill_name=skill).first()
            if not user_skill:
                user_skill = UserSkill(
                    user_id=user_id,
                    skill_name=skill,
                    skill_level='intermediate'
                )
                db.session.add(user_skill)
        
        try:
            db.session.commit()
        except Exception as e:
            db.session.rollback()
